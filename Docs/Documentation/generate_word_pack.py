"""
Generate Word (.docx) copies of manager documentation.

Docs only — does not change application code.

Run from hms-backend:

  python Docs/Documentation/generate_word_pack.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parent
MD_FILES = [
    ROOT / "00-HMS-Overview.md",
    ROOT / "01-Roles-and-Permissions.md",
    ROOT / "02-End-to-End-Patient-Journey.md",
    *sorted((ROOT / "Roles").glob("*.md")),
]


def _add_runs(paragraph, text: str) -> None:
    """Minimal markdown: **bold** segments."""
    parts = text.split("**")
    for i, part in enumerate(parts):
        run = paragraph.add_run(part)
        run.bold = i % 2 == 1
        run.font.size = Pt(11)


def md_to_docx(md_path: Path, docx_path: Path) -> None:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = md_path.read_text(encoding="utf-8").splitlines()
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if not line.strip():
            i += 1
            continue

        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
        elif line.startswith("```"):
            i += 1
            block: list[str] = []
            while i < len(lines) and not lines[i].startswith("```"):
                block.append(lines[i])
                i += 1
            p = doc.add_paragraph("\n".join(block))
            for run in p.runs:
                run.font.name = "Consolas"
                run.font.size = Pt(9)
        elif line.startswith("|") and "|" in line[1:]:
            rows: list[list[str]] = []
            while i < len(lines) and lines[i].startswith("|"):
                raw = lines[i].strip()
                cells = [c.strip() for c in raw.strip("|").split("|")]
                # skip markdown separator rows like |---|---|
                if all(set(c) <= set("-: ") for c in cells):
                    i += 1
                    continue
                rows.append(cells)
                i += 1
            if rows:
                cols = max(len(r) for r in rows)
                table = doc.add_table(rows=len(rows), cols=cols)
                table.style = "Table Grid"
                for r_idx, row in enumerate(rows):
                    for c_idx in range(cols):
                        cell_text = row[c_idx] if c_idx < len(row) else ""
                        table.rows[r_idx].cells[c_idx].text = cell_text
            continue
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, line[2:].strip())
        elif line.startswith("> "):
            p = doc.add_paragraph()
            _add_runs(p, line[2:].strip())
            for run in p.runs:
                run.italic = True
        elif set(line.strip()) <= {"-", "="} and len(line.strip()) >= 3:
            pass
        else:
            p = doc.add_paragraph()
            _add_runs(p, line)

        i += 1

    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)


def main() -> None:
    out_dir = ROOT / "Word"
    out_dir.mkdir(exist_ok=True)
    roles_out = out_dir / "Roles"
    roles_out.mkdir(exist_ok=True)

    for md in MD_FILES:
        if not md.exists():
            print(f"SKIP missing: {md}")
            continue
        if md.parent.name == "Roles":
            target = roles_out / f"{md.stem}.docx"
        else:
            target = out_dir / f"{md.stem}.docx"
        md_to_docx(md, target)
        print(f"Wrote {target.relative_to(ROOT)}")

    print("Done. Open Docs/Documentation/Word/ in Microsoft Word.")


if __name__ == "__main__":
    main()
