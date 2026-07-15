"""Generator: Doctor Profile & Notifications — Frontend Developer Guide (DOCX)."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

OUT = Path(__file__).resolve().parent / "doctor-profile-notifications-frontend-guide.docx"


def _set_run(run, *, bold=False, size=11, color=None, italic=False, font="Calibri"):
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    if color is not None:
        run.font.color.rgb = color


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        if level == 1:
            _set_run(run, bold=True, size=16, color=RGBColor(20, 60, 120))
        elif level == 2:
            _set_run(run, bold=True, size=13, color=RGBColor(30, 50, 90))
        else:
            _set_run(run, bold=True, size=11, color=RGBColor(40, 40, 40))
    return p


def add_para(doc, text, *, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    _set_run(run, bold=bold, italic=italic, size=size)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    run = p.add_run(text)
    _set_run(run, size=11)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text.strip("\n"))
    _set_run(run, size=9, font="Consolas")
    # Light grey background via shading
    shd = p._p.get_or_add_pPr()
    shd_elem = shd.makeelement(
        qn("w:shd"),
        {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): "F2F2F2",
        },
    )
    shd.append(shd_elem)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(h)
        _set_run(run, bold=True, size=9)
    for r_idx, row in enumerate(rows):
        cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ""
            run = cells[c_idx].paragraphs[0].add_run(str(val))
            _set_run(run, size=9)
    if col_widths:
        for row in table.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return table


def build_docx() -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    # ---- Cover ----
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Doctor Profile & Notifications")
    _set_run(run, bold=True, size=22, color=RGBColor(20, 60, 120))

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Frontend Developer Guide")
    _set_run(run, bold=True, size=16, color=RGBColor(20, 60, 120))

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(
        "Version: 2.0  |  Date: July 2026  |  Audience: Frontend developers\n"
        "Backend: HM-System (FastAPI)  |  Auth: JWT Bearer token\n"
        "Modules: /doctor/profile  +  /doctor/notifications"
    )
    _set_run(run, size=10, color=RGBColor(80, 80, 80))

    add_para(
        doc,
        "This document is the source of truth for implementing Doctor Profile and "
        "Doctor Notifications in the frontend. It matches the current nested API "
        "schemas (not the older flat field layout).",
        italic=True,
        size=10,
    )

    # =========================================================
    add_heading(doc, "1. Purpose & Scope", 1)

    add_para(
        doc,
        "Use this guide for endpoints, request/response shapes, validation, "
        "permissions, field ownership, and suggested UI behaviour.",
    )

    add_heading(doc, "What this covers", 2)
    for t in [
        "Doctor self-service profile: view, update, upload/delete profile image",
        "Which fields doctors can edit vs which are admin-only (read-only in UI)",
        "Nested payload shapes (address, emergency_contact, department, role, shift)",
        "Profile completion boolean + completion percentage",
        "In-app notifications: list, unread badge, mark one/all as read",
        "Notification types, priorities, filters, and deep-link hints",
        "TypeScript types and frontend checklist",
    ]:
        add_bullet(doc, t)

    add_heading(doc, "What this does NOT cover", 2)
    for t in [
        "Admin staff APIs (/users) except where they own doctor fields",
        "Push notifications, WebSocket, or SSE (not implemented — use polling)",
        "Delete notification or notification preferences (not implemented)",
    ]:
        add_bullet(doc, t)

    add_heading(doc, "Auth header (all endpoints)", 2)
    add_code(doc, "Authorization: Bearer <access_token>")
    add_para(
        doc,
        "Login via POST /auth/login. JWT must include the required permission. "
        "Inactive accounts → 403. Soft-deleted / missing users → 404.",
    )

    # =========================================================
    add_heading(doc, "2. Module A — Doctor Profile", 1)
    add_para(
        doc,
        "Prefix: /doctor  |  Router: Routers/doctor_profile_router.py  |  Tag: Doctor Profile",
    )
    add_para(
        doc,
        'Only users with role name "doctor" can use these endpoints. Admins manage '
        "license, fee, specialization, department, employee_id, joining_date, and "
        "shift via /users — not via /doctor/profile.",
    )

    add_heading(doc, "2.1 Endpoint summary", 2)
    add_table(
        doc,
        ["Method + Path", "Status", "Permission", "Purpose"],
        [
            ["GET /doctor/profile", "200", "doctor_profile:view", "Load full profile"],
            ["PUT /doctor/profile", "200", "doctor_profile:update", "Update editable fields"],
            ["POST /doctor/profile/image", "200", "doctor_profile:upload_image", "Upload avatar"],
            ["DELETE /doctor/profile/image", "200", "doctor_profile:delete_image", "Remove avatar"],
        ],
        col_widths=[2.2, 0.7, 2.0, 1.8],
    )

    # ---- GET ----
    add_heading(doc, "2.2 GET /doctor/profile", 2)
    add_para(doc, "Request body: none. Returns DoctorProfileResponse (nested objects).")
    add_code(
        doc,
        """{
  "user_id": 12,
  "first_name": "Asha",
  "last_name": "Mehta",
  "email": "asha.mehta@hospital.com",
  "phone": "9876543210",
  "phone_code": "+91",
  "address": {
    "line": "12 MG Road",
    "city": "Pune",
    "state": "Maharashtra"
  },
  "date_of_birth": "1985-04-12",
  "gender": 2,
  "emergency_contact": {
    "name": "Raj Mehta",
    "phone": "9123456780"
  },
  "department": { "id": 3, "name": "Cardiology" },
  "role": { "id": 5, "name": "doctor" },
  "specialization": "Interventional Cardiology",
  "qualification": "MD, DM Cardiology",
  "medical_license_number": "MH-MED-12345",
  "employee_id": "DOC-0012",
  "experience_years": 12,
  "joining_date": "2018-06-01",
  "consultation_fee": 800.0,
  "bio": "Consultant cardiologist with 12 years experience.",
  "languages": ["English", "Hindi", "Marathi"],
  "shift": {
    "name": "Morning",
    "start_time": "09:00",
    "end_time": "13:00"
  },
  "profile_image_url": "/uploads/doctor_image/a1b2c3d4.jpg",
  "is_profile_completed": true,
  "profile_completion_percentage": 86,
  "is_active": true,
  "last_login": "2026-07-15T09:00:00+05:30",
  "created_at": "2026-01-10T10:00:00+05:30",
  "updated_at": "2026-07-14T18:20:00+05:30"
}""",
    )

    add_para(doc, "Response field reference:", bold=True)
    add_table(
        doc,
        ["Field", "Type", "Notes"],
        [
            ["user_id", "number", "Doctor user id"],
            ["first_name", "string", "Read-only for doctor"],
            ["last_name", "string | null", "Read-only for doctor"],
            ["email", "string (email)", "Read-only for doctor"],
            ["phone", "string | null", "Editable"],
            ["phone_code", "string | null", "Editable (e.g. +91)"],
            ["address", "object", "{ line, city, state } — editable"],
            ["date_of_birth", "string | null", "YYYY-MM-DD, editable"],
            ["gender", "number | null", "1–4 — see Gender codes"],
            ["emergency_contact", "object", "{ name, phone } — editable"],
            ["department", "object | null", "{ id, name } — read-only"],
            ["role", "object | null", "{ id, name } — read-only"],
            ["specialization", "string | null", "Admin-owned — read-only"],
            ["qualification", "string | null", "Editable"],
            ["medical_license_number", "string | null", "Admin-owned — read-only"],
            ["employee_id", "string | null", "Admin-owned — read-only"],
            ["experience_years", "number | null", "Editable, 0–60"],
            ["joining_date", "string | null", "Admin-owned — read-only"],
            ["consultation_fee", "number | null", "Admin-owned — read-only"],
            ["bio", "string | null", "Editable"],
            ["languages", "string[]", "Editable, default []"],
            ["shift", "object | null", "{ name, start_time, end_time } — admin — read-only"],
            ["profile_image_url", "string | null", "Public path under /uploads"],
            ["is_profile_completed", "boolean", "Computed server-side (strict)"],
            ["profile_completion_percentage", "number", "0–100 progress bar value"],
            ["is_active", "boolean", "Account status"],
            ["last_login", "datetime | null", "ISO-8601"],
            ["created_at / updated_at", "datetime | null", "Profile timestamps"],
        ],
    )

    add_heading(doc, "Gender codes (integer, not label)", 3)
    for t in ["1 = Male", "2 = Female", "3 = Other", "4 = Prefer not to say"]:
        add_bullet(doc, t)
    add_para(doc, "API returns the integer. Map to labels in the UI.")

    add_heading(doc, "Errors (GET)", 3)
    for t in [
        "401 — Missing/invalid JWT",
        "403 — Missing permission, inactive account, or role is not doctor",
        '404 — User/profile not found ("Doctor profile not found. Contact admin.")',
    ]:
        add_bullet(doc, t)

    # ---- PUT ----
    add_heading(doc, "2.3 PUT /doctor/profile", 2)
    add_para(
        doc,
        'Content-Type: application/json. Schema: DoctorProfileUpdate with extra="forbid". '
        "Unknown or admin-only fields cause 422. At least one field required; empty body → 400.",
    )
    add_para(doc, "All fields optional; send only what changed. Nested objects are partial (merge by key):")
    add_code(
        doc,
        """{
  "qualification": "MD, DM Cardiology",
  "experience_years": 12,
  "bio": "Consultant cardiologist...",
  "languages": ["English", "Hindi", "Marathi"],
  "phone": "9876543210",
  "phone_code": "+91",
  "address": {
    "line": "12 MG Road",
    "city": "Pune",
    "state": "Maharashtra"
  },
  "date_of_birth": "1985-04-12",
  "gender": 2,
  "emergency_contact": {
    "name": "Raj Mehta",
    "phone": "9123456780"
  }
}""",
    )

    add_para(doc, "Validation rules:", bold=True)
    for t in [
        "qualification: max 255 chars",
        "experience_years: integer, >= 0 and <= 60",
        "phone / emergency_contact.phone: max 20 chars",
        "phone_code: max 10 chars",
        "address.city / address.state: max 100 chars",
        "emergency_contact.name: max 120 chars",
        "gender: integer 1–4",
        "languages: trimmed; empties dropped; case-insensitive dedupe",
        "Do NOT send flat city/state/address strings — use nested address object",
        "Do NOT send: first_name, last_name, email, department, role, specialization, "
        "medical_license_number, employee_id, joining_date, consultation_fee, shift, "
        "profile_image_url, is_profile_completed, profile_completion_percentage",
    ]:
        add_bullet(doc, t)

    add_para(
        doc,
        "Response: same DoctorProfileResponse as GET. "
        "Side effect: is_profile_completed and profile_completion_percentage recomputed.",
    )

    add_heading(doc, "Profile completion rules", 3)
    add_para(
        doc,
        "is_profile_completed is true only when ALL of these are set: "
        "qualification (truthy), experience_years (not null — 0 counts), bio (truthy). "
        "Image, languages, phone, license, and fee are NOT required for this boolean.",
    )
    add_para(
        doc,
        "profile_completion_percentage (0–100) is a separate progress metric based on "
        "14 checks: phone, phone_code, address line, city, state, DOB, gender, "
        "emergency name, emergency phone, qualification, experience_years, bio, "
        "languages (non-empty), profile image. Use this for a progress bar; use "
        "is_profile_completed for the incomplete banner / gate.",
    )

    add_heading(doc, "UI field ownership (important)", 3)
    for t in [
        "EDITABLE by doctor: qualification, experience_years, bio, languages, phone, "
        "phone_code, address.{line,city,state}, date_of_birth, gender, "
        "emergency_contact.{name,phone}, profile image",
        "READ-ONLY display: first_name, last_name, email, department, role, "
        "specialization, medical_license_number, employee_id, joining_date, "
        "consultation_fee, shift, is_active, is_profile_completed, "
        "profile_completion_percentage",
        "Show a short note that license / fee / specialization / department / shift "
        "are managed by admin",
    ]:
        add_bullet(doc, t)

    # ---- Image ----
    add_heading(doc, "2.4 POST /doctor/profile/image", 2)
    add_para(doc, "Content-Type: multipart/form-data. Field name MUST be: file")
    for t in [
        "Allowed extensions: .jpg, .jpeg, .png, .webp",
        "Max size: 5 MB",
        "Empty file rejected",
        "Replaces previous image (old file deleted on server)",
        "Image upload does NOT flip is_profile_completed by itself (but does raise "
        "profile_completion_percentage if image was missing)",
    ]:
        add_bullet(doc, t)

    add_code(
        doc,
        """// Example (fetch)
const form = new FormData();
form.append("file", selectedFile);
await fetch(`${API_BASE}/doctor/profile/image`, {
  method: "POST",
  headers: { Authorization: `Bearer ${token}` },
  // Do NOT set Content-Type manually — browser sets multipart boundary
  body: form
});

// Success response
{
  "message": "Profile image uploaded successfully",
  "profile_image_url": "/uploads/doctor_image/<uuid>.jpg"
}""",
    )
    add_para(
        doc,
        "Display image as: {API_BASE}{profile_image_url} "
        "e.g. http://localhost:8000/uploads/doctor_image/....jpg "
        "(static files mounted at /uploads).",
    )
    for t in [
        "400 — Missing filename, wrong type, empty, > 5 MB, invalid path",
        "404 — Profile not found",
        "500 — Failed to save file",
    ]:
        add_bullet(doc, t)

    add_heading(doc, "2.5 DELETE /doctor/profile/image", 2)
    add_para(doc, "Request body: none.")
    add_code(
        doc,
        """{
  "message": "Profile image deleted successfully",
  "profile_image_url": null
}""",
    )
    add_bullet(doc, '404 — "No profile image to delete" or profile missing')

    add_heading(doc, "2.6 Suggested Profile UI", 2)
    for t in [
        "Profile Overview: avatar, name, department.name, specialization, fee, "
        "shift, completion badge + progress bar",
        "Edit Professional: qualification, experience_years, bio, languages (tag input)",
        "Edit Contact: phone + phone_code, address fields, DOB, gender select, "
        "emergency name + phone",
        "Avatar: enforce type + 5 MB client-side before upload",
        "On load: GET once; spinner; if 404 show Contact admin",
        "On save: PUT only dirty fields (and nested objects with only dirty keys); "
        "refresh form from response",
        "Show incomplete banner when is_profile_completed === false",
        "Bind progress bar to profile_completion_percentage",
    ]:
        add_bullet(doc, t)

    add_heading(doc, "2.7 TypeScript types (Profile)", 2)
    add_code(
        doc,
        """export type GenderCode = 1 | 2 | 3 | 4;

export interface AddressInfo {
  line: string | null;
  city: string | null;
  state: string | null;
}

export interface EmergencyContactInfo {
  name: string | null;
  phone: string | null;
}

export interface DepartmentInfo { id: number; name: string; }
export interface RoleInfo { id: number; name: string; }
export interface ShiftInfo {
  name: string | null;
  start_time: string | null; // "HH:MM"
  end_time: string | null;
}

export interface DoctorProfile {
  user_id: number;
  first_name: string;
  last_name: string | null;
  email: string;
  phone: string | null;
  phone_code: string | null;
  address: AddressInfo;
  date_of_birth: string | null; // YYYY-MM-DD
  gender: GenderCode | null;
  emergency_contact: EmergencyContactInfo;
  department: DepartmentInfo | null;
  role: RoleInfo | null;
  specialization: string | null;
  qualification: string | null;
  medical_license_number: string | null;
  employee_id: string | null;
  experience_years: number | null;
  joining_date: string | null;
  consultation_fee: number | null;
  bio: string | null;
  languages: string[];
  shift: ShiftInfo | null;
  profile_image_url: string | null;
  is_profile_completed: boolean;
  profile_completion_percentage: number;
  is_active: boolean;
  last_login: string | null;
  created_at: string | null;
  updated_at: string | null;
}

export interface DoctorProfileUpdate {
  qualification?: string | null;
  experience_years?: number | null; // 0-60
  bio?: string | null;
  languages?: string[] | null;
  phone?: string | null;
  phone_code?: string | null;
  address?: Partial<AddressInfo> | null;
  date_of_birth?: string | null;
  gender?: GenderCode | null;
  emergency_contact?: Partial<EmergencyContactInfo> | null;
}

export interface DoctorProfileImageResponse {
  message: string;
  profile_image_url: string | null;
}""",
    )

    # =========================================================
    add_heading(doc, "3. Module B — Doctor Notifications", 1)
    add_para(
        doc,
        "Prefix: /doctor/notifications  |  Router: Routers/doctor_notification_router.py  |  "
        "Tag: Doctor Notifications",
    )
    add_para(
        doc,
        "In-app notifications for the logged-in doctor only (scoped by user_id). "
        "No WebSocket/SSE/push — poll unread-count for the nav badge. "
        "Doctors cannot create notifications; they are produced by OPD, Lab, Nurse, Admin flows.",
    )

    add_heading(doc, "3.1 Endpoint summary", 2)
    add_table(
        doc,
        ["Method + Path", "Status", "Permission", "Purpose"],
        [
            ["GET /doctor/notifications", "200", "notifications:view", "Paginated list + filters"],
            ["GET .../unread-count", "200", "notifications:view", "Badge count"],
            ["PATCH .../{id}/read", "200", "notifications:update", "Mark one as read"],
            ["PATCH .../read-all", "200", "notifications:update", "Mark all as read"],
        ],
        col_widths=[2.6, 0.7, 1.8, 1.8],
    )

    add_heading(doc, "Not implemented (do not build UI yet)", 3)
    for t in [
        "Delete notification",
        "Preferences / mute / channels",
        "Push (FCM/APNs) or real-time WebSocket/SSE",
        "Doctor-facing create API",
    ]:
        add_bullet(doc, t)

    add_heading(doc, "3.2 GET /doctor/notifications", 2)
    add_para(doc, "Query parameters:")
    add_table(
        doc,
        ["Param", "Type", "Rules"],
        [
            ["page", "int", ">= 1, default 1"],
            ["limit", "int", "1–100, default 20"],
            ["search", "string", "min length 1; ILIKE on title, message, created_by_name"],
            ["is_read", "bool", "Optional filter true/false"],
            ["source_module", "enum", "See SourceModule"],
            ["notification_type", "enum", "See NotificationType"],
            ["start_date", "date", "YYYY-MM-DD, inclusive (Asia/Kolkata)"],
            ["end_date", "date", "YYYY-MM-DD, inclusive (Asia/Kolkata)"],
        ],
    )
    add_para(doc, "Sort order: priority CRITICAL > HIGH > NORMAL, then created_at descending.")
    add_code(
        doc,
        """GET /doctor/notifications?page=1&limit=20&is_read=false

{
  "total": 42,
  "page": 1,
  "limit": 20,
  "items": [
    {
      "id": 101,
      "user_id": 12,
      "title": "Lab Report Ready",
      "message": "CBC — Ramesh Patil",
      "notification_type": "LAB_REPORT_READY",
      "priority": "HIGH",
      "source_module": "LAB",
      "reference_type": "LAB_ORDER",
      "reference_id": 55,
      "created_by": 8,
      "created_by_name": "Lab Tech One",
      "is_read": false,
      "read_at": null,
      "created_at": "2026-07-13T09:15:00+05:30"
    }
  ]
}""",
    )

    add_para(doc, "NotificationResponse fields:", bold=True)
    add_table(
        doc,
        ["Field", "Type", "Notes"],
        [
            ["id", "number", "Primary key"],
            ["user_id", "number", "Recipient doctor user id"],
            ["title", "string", "Max 255"],
            ["message", "string | null", "May contain \\n — render multiline"],
            ["notification_type", "enum string", "See types table"],
            ["priority", "enum string", "NORMAL | HIGH | CRITICAL"],
            ["source_module", "enum string", "Producer module"],
            ["reference_type", "enum string", "Deep-link entity type"],
            ["reference_id", "number", "Deep-link entity id"],
            ["created_by", "number | null", "Actor user id"],
            ["created_by_name", "string | null", "Display name snapshot"],
            ["is_read", "boolean", ""],
            ["read_at", "datetime | null", "ISO-8601 with timezone"],
            ["created_at", "datetime", "ISO-8601 with timezone"],
        ],
    )

    add_heading(doc, "3.3 GET /doctor/notifications/unread-count", 2)
    add_code(doc, '{ "count": 5 }')
    add_para(
        doc,
        "Use for the nav bell badge. Poll every 15–30 seconds (or on focus/visibility change). "
        "No real-time channel exists.",
    )

    add_heading(doc, "3.4 PATCH /doctor/notifications/{notification_id}/read", 2)
    add_para(doc, "Body: none. Returns full NotificationResponse. Idempotent if already read.")
    for t in [
        "Sets is_read=true and read_at=now (Asia/Kolkata) on first mark",
        '404 — "Notification not found" (wrong id or not owned by current user)',
    ]:
        add_bullet(doc, t)

    add_heading(doc, "3.5 PATCH /doctor/notifications/read-all", 2)
    add_para(doc, "Body: none.")
    add_code(doc, '{ "message": "All notifications marked as read" }')
    add_para(doc, "After success, set local unread count to 0 and refresh the list if open.")

    add_heading(doc, "3.6 Enums", 2)
    add_para(doc, "NotificationPriority: NORMAL | HIGH | CRITICAL")
    add_para(doc, "Default priority by type:")
    for t in [
        "NEW_APPOINTMENT → NORMAL",
        "LAB_REPORT_READY, APPOINTMENT_CANCELLED, APPOINTMENT_RESCHEDULED, "
        "ADMIN_UPDATE, HANDOVER_TAKEN_OVER, SHIFT_UPDATED → HIGH",
        "EMERGENCY_ALERT → CRITICAL",
    ]:
        add_bullet(doc, t)

    add_para(doc, "NotificationType — currently produced for doctors:", bold=True)
    add_table(
        doc,
        ["Type", "When created", "Typical deep-link"],
        [
            [
                "NEW_APPOINTMENT",
                "Paid patient added to doctor queue (check-in)",
                "APPOINTMENT + appointment id",
            ],
            [
                "APPOINTMENT_CANCELLED",
                "Appointment cancelled AND visit paid",
                "APPOINTMENT + appointment id",
            ],
            [
                "APPOINTMENT_RESCHEDULED",
                "scheduled_at changed, not cancelled, paid",
                "APPOINTMENT + appointment id",
            ],
            [
                "LAB_REPORT_READY",
                "Lab report uploaded / first file report created",
                "LAB_ORDER + order id",
            ],
            [
                "ADMIN_UPDATE",
                "Admin dept change / deactivate / delete / profile admin edits",
                "USER + user id",
            ],
        ],
    )

    add_para(doc, "Reserved / other enum values (exist; may appear in filters or future flows):")
    for t in [
        "PATIENT_CHECKED_IN, LAB_REPORT_UPDATED, PRESCRIPTION_CREATED, PRESCRIPTION_UPDATED",
        "EMERGENCY_ALERT (mainly nurse flows; doctors may still receive alerts in some cases)",
        "HANDOVER_TAKEN_OVER, SHIFT_UPDATED (primarily nurse; include in TS union for shared types)",
    ]:
        add_bullet(doc, t)

    add_para(
        doc,
        "SourceModule: OPD_BILLING | LAB | NURSE | ADMIN (active for doctors). "
        "Also defined: RECEPTIONIST, PHARMACY, SYSTEM",
    )
    add_para(doc, "ReferenceType (deep-link):")
    for t in [
        "APPOINTMENT + reference_id → appointment / queue / calendar screen",
        "LAB_ORDER + reference_id → lab order detail",
        "PATIENT + reference_id → patient chart / emergency context",
        "USER + reference_id → profile / account notice (admin updates)",
        "Also defined: PRESCRIPTION, BILL, SCHEDULE, LEAVE, HANDOVER, ALERT",
    ]:
        add_bullet(doc, t)

    add_heading(doc, "3.7 Notification trigger details (UI copy)", 2)
    for t in [
        'NEW_APPOINTMENT title: "Paid Appointment Confirmed" — patient, time, token',
        'APPOINTMENT_CANCELLED title: "Appointment Cancelled" — patient + time',
        'APPOINTMENT_RESCHEDULED title: "Appointment Rescheduled" — patient + new time',
        'LAB_REPORT_READY title: "Lab Report Ready" — "{test} — {patient}"',
        'ADMIN_UPDATE — e.g. department reassigned, account disabled by admin',
        "Unpaid appointments do NOT notify on cancel/reschedule",
        "Replacing an existing lab file report does NOT create another notification",
        "Message often contains newlines (\\n) — render with white-space: pre-line",
    ]:
        add_bullet(doc, t)

    add_heading(doc, "3.8 Suggested Notifications UI", 2)
    for t in [
        "Nav: bell icon + unread badge from unread-count (poll)",
        "Inbox list: title, message (multiline), priority chip, type, relative time, unread style",
        "Filters: Unread / All, type chips, date range, search box",
        "Actions: click row → mark read + navigate via reference_type/reference_id",
        'Toolbar: "Mark all as read"',
        "Empty state when total === 0",
        "Visual priority: CRITICAL red, HIGH amber, NORMAL default",
        "Do not invent delete/archive/preferences UI until backend supports it",
    ]:
        add_bullet(doc, t)

    add_heading(doc, "3.9 TypeScript types (Notifications)", 2)
    add_code(
        doc,
        """export type NotificationPriority = "NORMAL" | "HIGH" | "CRITICAL";

export type NotificationType =
  | "NEW_APPOINTMENT"
  | "APPOINTMENT_CANCELLED"
  | "APPOINTMENT_RESCHEDULED"
  | "PATIENT_CHECKED_IN"
  | "LAB_REPORT_READY"
  | "LAB_REPORT_UPDATED"
  | "PRESCRIPTION_CREATED"
  | "PRESCRIPTION_UPDATED"
  | "EMERGENCY_ALERT"
  | "ADMIN_UPDATE"
  | "HANDOVER_TAKEN_OVER"
  | "SHIFT_UPDATED";

export type SourceModule =
  | "OPD_BILLING" | "LAB" | "RECEPTIONIST" | "NURSE"
  | "PHARMACY" | "ADMIN" | "SYSTEM";

export type ReferenceType =
  | "APPOINTMENT" | "LAB_ORDER" | "PRESCRIPTION" | "BILL"
  | "PATIENT" | "USER" | "SCHEDULE" | "LEAVE"
  | "HANDOVER" | "ALERT";

export interface Notification {
  id: number;
  user_id: number;
  title: string;
  message: string | null;
  notification_type: NotificationType;
  priority: NotificationPriority;
  source_module: SourceModule;
  reference_type: ReferenceType;
  reference_id: number;
  created_by: number | null;
  created_by_name: string | null;
  is_read: boolean;
  read_at: string | null;
  created_at: string;
}

export interface NotificationListResponse {
  total: number;
  page: number;
  limit: number;
  items: Notification[];
}

export interface UnreadCountResponse {
  count: number;
}""",
    )

    # =========================================================
    add_heading(doc, "4. Error Handling Cheat Sheet", 1)
    add_table(
        doc,
        ["HTTP", "When / Frontend action"],
        [
            ["400", "Empty profile PUT; bad image upload — show toast from detail"],
            ["401", "Redirect to login; clear token"],
            ["403", "Permission denied / inactive / not doctor — show access denied"],
            ["404", "Profile missing (contact admin); notification not found; no image"],
            ["422", "Validation / forbidden extra fields — show field errors"],
            ["500", "Image save failure — retry later"],
        ],
        col_widths=[0.8, 5.5],
    )
    add_para(
        doc,
        'FastAPI usually returns { "detail": "..." } or a validation array for 422. '
        "Surface detail in toasts; map 422 to form fields when possible.",
    )

    # =========================================================
    add_heading(doc, "5. Frontend Implementation Checklist", 1)
    add_heading(doc, "Profile", 2)
    for t in [
        "[ ] Route/page: Doctor Profile under authenticated doctor layout",
        "[ ] GET /doctor/profile on mount",
        "[ ] Separate editable vs read-only sections visually",
        "[ ] Bind nested address + emergency_contact correctly (not flat strings)",
        "[ ] Gender select mapped to 1–4",
        "[ ] Languages as chip/tag input",
        "[ ] PUT with only changed fields; nested partial objects OK",
        '[ ] Avatar upload via FormData field "file"; preview + 5 MB client check',
        "[ ] Avatar delete + clear preview",
        "[ ] Show is_profile_completed banner until true",
        "[ ] Progress bar from profile_completion_percentage",
        "[ ] Image URL = API_BASE + profile_image_url",
        "[ ] Display shift, employee_id, joining_date as read-only if present",
    ]:
        add_bullet(doc, t)

    add_heading(doc, "Notifications", 2)
    for t in [
        "[ ] Bell in doctor header with unread badge",
        "[ ] Poll GET /doctor/notifications/unread-count (15–30s)",
        "[ ] Notifications page/drawer with pagination",
        "[ ] Filters: is_read, notification_type, date range, search",
        "[ ] Mark one read on open; Mark all read button",
        "[ ] Deep-link by reference_type + reference_id",
        "[ ] Render message with line breaks; priority colours",
        "[ ] Optimistic unread decrement after mark-read",
    ]:
        add_bullet(doc, t)

    # =========================================================
    add_heading(doc, "6. Permissions Required (Doctor role)", 1)
    for t in [
        "doctor_profile:view",
        "doctor_profile:update",
        "doctor_profile:upload_image",
        "doctor_profile:delete_image",
        "notifications:view",
        "notifications:update",
    ]:
        add_bullet(doc, t)
    add_para(
        doc,
        "These are seeded for the doctor role. Permissions are embedded in the JWT at login. "
        "Gate UI routes/buttons on permissions if your app checks them client-side.",
    )

    # =========================================================
    add_heading(doc, "7. Quick API Reference", 1)
    add_code(
        doc,
        """# Profile
GET    /doctor/profile
PUT    /doctor/profile
POST   /doctor/profile/image          (multipart: file)
DELETE /doctor/profile/image

# Notifications
GET    /doctor/notifications
GET    /doctor/notifications/unread-count
PATCH  /doctor/notifications/{id}/read
PATCH  /doctor/notifications/read-all

# Auth
Authorization: Bearer <token>
OpenAPI: /docs when backend is running""",
    )

    # =========================================================
    add_heading(doc, "8. Source Files (Backend)", 1)
    for t in [
        "Routers/doctor_profile_router.py",
        "Schemas/doctor_profile_schema.py",
        "Services/doctor_profile_service.py",
        "Models/doctor_profile.py",
        "Routers/doctor_notification_router.py",
        "Schemas/notification_schema.py",
        "Enums/notification.py",
        "Models/notification.py",
        "Services/notification_service.py",
        "seed.py (permissions)",
    ]:
        add_bullet(doc, t)

    add_para(
        doc,
        "End of document. For live schema details, run the backend and open /docs. "
        "Regenerate this file with: python Docs/generate_doctor_profile_notification_docx.py",
        italic=True,
        size=9,
    )

    doc.save(str(OUT))
    return OUT


if __name__ == "__main__":
    path = build_docx()
    print(f"Wrote {path}")
